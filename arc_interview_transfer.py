'''
Program to aid ARC interview data transfer.

Author: Jorie Sieck
Website: www.joriesieck.com
'''
import docx
import openpyxl
import os

file_names = os.listdir('1-segmented interview transcripts') # gets the names of the files to read from
 
def get_text(file):
    '''gets the interview data and stores it in a list, with each segment as a string.'''
    full_text = []
    str_to_append = ''
    i = 0
    while len(file.paragraphs) > i: # run through each paragraph in the file
        if '-seg' in str(file.paragraphs[i].text) or file.paragraphs[i].text == 'END': # check for markers to split the text
            if str_to_append != '': # this just prevents it from adding the empty string when it hits the first marker
                full_text.append(str_to_append) # append the string up to but not including the marker to the list
            str_to_append = str(file.paragraphs[i].text) # reset the string so it only contains the marker
        else:
            str_to_append += str(file.paragraphs[i].text) # add the current paragraph to the string
        i += 1 # increment the index
    return full_text
 
def write_text(w_file,w_filename,x_file,x_sheet,x_filename):
    '''takes the listified interview data from get_text and writes to an excel spreadsheet.'''
    doc_text = get_text(w_file)
    for i in range(len(doc_text)):
        line_num = 1
        while x_sheet['A' + str(line_num)].value != None:
            line_num += 1
        try:
            seg = int(doc_text[i][4] + doc_text[i][5])
            try:
                task = int(doc_text[i][11] + doc_text[i][12])
            except:
                task = int(doc_text[i][11])
        except:
            seg = int(doc_text[i][4])
            try:
                task = int(doc_text[i][10] + doc_text[i][11])
            except:
                task = int(doc_text[i][10])
        x_sheet['A' + str(line_num)].value = w_filename
        x_sheet['B' + str(line_num)].value = int(w_filename[0])
        x_sheet['C' + str(line_num)].value = task
        x_sheet['D' + str(line_num)].value = seg
        x_sheet['E' + str(line_num)].value = doc_text[i]
        seg += 1
        i += 1
    x_file.save(x_filename)

def main():
    xl_filename = 'interview data.xlsx'
    xl_file = openpyxl.load_workbook(xl_filename) # open the excel file that the data will be written to
    xl_sheet = xl_file['Sheet1'] # open the specific sheet that the data will be written to
    for i in range(len(file_names)):
        word_filename = file_names[i]
        this_file_dir = os.path.dirname(os.path.realpath('__file__'))
        word_file_dir = os.path.join(this_file_dir, '1-segmented interview transcripts/' + word_filename)
        word_file = docx.Document(word_file_dir) # open the word document that the data comes from
        write_text(word_file,word_filename,xl_file,xl_sheet,xl_filename)

main()